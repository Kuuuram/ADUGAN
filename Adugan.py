import streamlit as st
from streamlit_option_menu import option_menu
from datetime import date
from datetime import time
from datetime import datetime
from docx import Document
from docx.shared import Inches
import pathlib

#Page set
st.set_page_config(page_title="ADUGAN", page_icon= ":speaking_head_in_silhouette:", layout="centered")

#Fungsi simpan ke dokumen bagian Laporan
def save_to_docx(data, image_path, filename):
    doc = Document()
    doc.add_heading('Hasil Survei', 0)
    
    for key, value in data.items():
        doc.add_heading(key, level=1)
        doc.add_paragraph(str(value))
    
    if image_path:
        doc.add_heading('Lampiran Foto', level=1)
        doc.add_picture(image_path, width=Inches(4.0))
        doc.save(filename)
    else:
        doc.save(filename)

#Fungsi simpan ke dokumen bagian Aspirasi
def save_to_docs(data, filename):
    doc = Document()
    doc.add_heading('Hasil Survei', 0)
    
    for key, value in data.items():
        doc.add_heading(key, level=1)
        doc.add_paragraph(str(value))
        doc.save(filename)

#Fungsi memanggil CSS
def load_css(file_path):
    with open(file_path) as f:
        st.markdown(f"<style>{f.read()}<style>", unsafe_allow_html=True)
css_path = pathlib.Path("C:/Users/iwaku/OneDrive/Documents/Python/Stel.css")
load_css(css_path)


#Deklarasi session state untuk tombol kirim
if 'submitted' not in st.session_state:
        st.session_state.submitted = False

#Navigasi bar
selected = option_menu(
    menu_title= "Mau mengadu apa?",
    options=["LAPORAN","ASPIRASI"],
    orientation="horizontal"
)
#Jika memilih bar LAPORAN
if selected == "LAPORAN":
    st.title(f"BUAT {selected}")
    Nama = st.text_input("Nama anda")
    Tanggal = st.date_input("Pilih tanggal:", value=date.today())
    Waktu = st.time_input("Waktu kejadian", value=time(12, 0))
    Korban = st.text_input("Nama korban")
    Pelaku = st.text_input("Nama pelaku")
    Tempat = st.text_input("Tempat kejadian")
    Isi = st.text_area("Isi laporan")
    uploaded_file = st.file_uploader("Unggah foto bukti", type=["jpg", "jpeg", "png"])
    image_path = None

    if uploaded_file is not None:
          image_path = f"uploaded_image.{uploaded_file.name.split('.')[-1]}"
          with open(image_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
            st.image(uploaded_file, caption="Foto yang diunggah", use_container_width=True)
    if st.button("Kirim", type="primary"):
        if not Nama:
            st.error("Isi nama anda/ketikkan anonim!")
        elif not Tanggal:
            st.error("Pilih tanggal kejadian!")
        elif not Waktu:
            st.error("Isi waktu kejadian!")
        elif not Korban:
            st.error("Isi nama korban!")
        elif not Pelaku:
            st.error("Isi nama pelaku!")
        elif not Tempat:
            st.error("Isi tempat kejadian!")
        elif not Isi:
            st.error("Isi laporan anda!")
        else:
            st.session_state.submitted = True
            st.write("Laporan berhasil terkirim!")
        survey_data = {
            "Nama pelapor" : Nama,
            "Tanggal kejadian" : Tanggal,
            "Waktu kejadian" : Waktu,
            "Nama korban" : Korban,
            "Nama pelaku" : Pelaku,
            "Tempat kejadian" : Tempat,
            "Isi laporan" : Isi
        }
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_filename = f"LAPORAN_{timestamp}.docx"
        save_to_docx(survey_data, image_path, doc_filename)

#Jika memilihh bar ASPIRASI
if selected == "ASPIRASI":
    st.title(f"SAMPAIKAN {selected}")
    Namas = st.text_input("Nama anda")
    Tujuan = st.text_input("Dijutukan untuk")
    Asp = st.text_area("Isi aspirasi")
    
    if st.button("Kirim", type="primary"):
        if not Namas:
            st.error("Isi nama anda/ketikkan anonim!")
        elif not Tujuan:
            st.error("Isi tujuan aspirasi!")
        elif not Asp:
            st.error("Isi aspirasi anda!")
        else:
            st.session_state.submitted = True
            st.write("Laporan berhasil terkirim!")

        survey_data = {
            "Nama pengirim" : Namas,
            "Ditujukan untuk" : Tujuan,
            "Isi aspirasi" : Asp,
        }
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_filenames = f"ASPIRASI_{timestamp}.docx"
        save_to_docs(survey_data, doc_filenames)    
